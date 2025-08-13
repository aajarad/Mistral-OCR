import os
import base64
from io import BytesIO

import streamlit as st
from dotenv import load_dotenv
from mistralai import Mistral
from docx import Document

# Load environment variables
load_dotenv()

# Helper: parse a 1-based page selection string like "1,3-5"
def parse_page_selection(s: str):
    """Return a set of 1-based page numbers, or None if empty/invalid."""
    if not s:
        return None
    s = s.strip()
    if not s:
        return None
    pages = set()
    for part in s.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            if a.strip().isdigit() and b.strip().isdigit():
                start = int(a)
                end = int(b)
                if start > end:
                    start, end = end, start
                pages.update(range(start, end + 1))
        elif part.isdigit():
            pages.add(int(part))
    return pages or None

st.set_page_config(page_title="Mistral OCR - Abdullah", page_icon="📖", layout="centered")

st.title("📖 Mistral OCR")
st.caption("Simple OCR UI by Abdullah – Convert your PDFs to Markdown using Mistral AI")

# Sidebar configuration
with st.sidebar:
    st.header("Settings")
    default_api_key = os.getenv("MISTRAL_API_KEY", "")
    api_key = st.text_input(
        "Mistral API Key",
        value=default_api_key,
        type="password",
        help="If empty, the app will try to use MISTRAL_API_KEY from your .env"
    )
    keep_images = st.checkbox("Include image base64 in response", value=False)
    pages_input = st.text_input(
        "Pages to include (optional)",
        value="",
        placeholder="Examples: 2 | 1,4 | 3-6 | 1,3-5,8",
        help=(
            "Select specific 1-based pages. Examples: 2 | 1,4 | 3-6 | 1,3-5,8. "
            "Leave empty to include ALL pages. Out-of-range pages are ignored."
        ),
    )

st.markdown("---")

uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"]) 

convert_btn = st.button("Convert to Markdown")

if convert_btn:
    if uploaded_file is None:
        st.error("Please upload a PDF file first.")
    else:
        key = api_key or os.getenv("MISTRAL_API_KEY", "")
        if not key:
            st.error("Mistral API Key is missing. Provide it in the sidebar or set MISTRAL_API_KEY in your .env.")
        else:
            try:
                # Read file bytes
                pdf_bytes = uploaded_file.read()
                b64 = base64.b64encode(pdf_bytes).decode("utf-8")

                # Initialize client
                client = Mistral(api_key=key)

                with st.spinner("Processing with Mistral OCR..."):
                    resp = client.ocr.process(
                        model="mistral-ocr-latest",
                        document={
                            "type": "document_url",
                            "document_url": f"data:application/pdf;base64,{b64}"
                        },
                        include_image_base64=keep_images,
                    )

                # Aggregate markdown with optional page filtering
                total_pages = len(resp.pages)
                sel_1based = parse_page_selection(pages_input)
                if sel_1based:
                    # keep only valid pages within bounds (convert to 0-based indices)
                    selected_idx = {p - 1 for p in sel_1based if 1 <= p <= total_pages}
                else:
                    selected_idx = None  # means all pages

                md_pages = []
                for page in resp.pages:
                    if selected_idx is None or page.index in selected_idx:
                        md_pages.append(f"## Page {page.index + 1}\n\n{page.markdown}")
                md_text = "\n\n".join(md_pages)

                if selected_idx is not None:
                    st.info(f"Showing {len(md_pages)} page(s) out of {total_pages} based on your selection.")
                else:
                    st.info(f"Showing all {total_pages} page(s).")

                st.success("Conversion complete!")
                st.subheader("Preview")
                st.markdown(md_text)

                # Download button
                base_name = os.path.splitext(uploaded_file.name)[0] or "output"
                md_bytes = md_text.encode("utf-8")
                st.download_button(
                    label="Download Markdown",
                    data=md_bytes,
                    file_name=f"{base_name}.md",
                    mime="text/markdown",
                )

                # Download as Word (.docx)
                doc = Document()
                # Simple mapping for headings; fallback to normal paragraphs
                for line in md_text.splitlines():
                    if line.startswith("###### "):
                        doc.add_heading(line[7:], level=6)
                    elif line.startswith("##### "):
                        doc.add_heading(line[6:], level=5)
                    elif line.startswith("#### "):
                        doc.add_heading(line[5:], level=4)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    else:
                        doc.add_paragraph(line)
                buf = BytesIO()
                doc.save(buf)
                st.download_button(
                    label="Download Word (.docx)",
                    data=buf.getvalue(),
                    file_name=f"{base_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"Failed to process the PDF: {e}")
                st.exception(e)

st.markdown("---")
st.caption("Built with ❤️ by Abdullah")
